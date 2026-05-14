"""
views.py  —  Two-step UX

Step A — GET  /           Show the upload form (image only at this stage)
         POST /           Receive image → call Claude → return SEO preview as JSON
                          The image is cached in Redis keyed by a session token.
                          No thotfy write yet.

Step B — POST /submit/    Receive partner name + product name + email
                          + the session token that points to the cached image + SEO content.
                          Kick off the Celery task. Redirect to /status/<task_id>/

Step C — GET  /status/<task_id>/
         Returns JSON for the frontend poller:
         { status, seo_title, meta_description, admin_url, error }
"""
from __future__ import annotations

import base64
import io
import logging
import uuid

from django.core.cache import cache
from django.contrib.auth.decorators import login_required
import docx
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render
from django.utils.decorators import method_decorator
from django.views import View
from django.views.decorators.csrf import ensure_csrf_cookie
from django.views.decorators.http import require_GET

from .services import ClaudeServiceError, generate_seo_content, get_hub_overview
from .tasks import cache_image, run_seo_pipeline

logger = logging.getLogger(__name__)

ALLOWED_TYPES = {
    "image/jpeg": "image/jpeg",
    "image/jpg":  "image/jpeg",
    "image/png":  "image/png",
    "image/webp": "image/webp",
}

SEO_CACHE_KEY  = "seo_result_{token}"
SEO_CACHE_TTL  = 60 * 30   # 30 min — user can take their time reviewing


@method_decorator(ensure_csrf_cookie, name='dispatch')
class GenerateView(View):
    """
    GET  / → render upload page
    POST / → receive image, call Claude, return SEO preview JSON
    """

    def get(self, request):
        return render(request, "seo_tool/index.html")

    def post(self, request):
        # ── Validate image ────────────────────────────────────────────────────
        image_file = request.FILES.get("product_image")
        if not image_file:
            return JsonResponse({"error": "No image uploaded."}, status=400)

        media_type = ALLOWED_TYPES.get(
            getattr(image_file, "content_type", "").lower()
        )
        if not media_type:
            return JsonResponse(
                {"error": "Unsupported format. Upload JPG, PNG or WebP."},
                status=400,
            )

        if image_file.size > 8 * 1024 * 1024:
            return JsonResponse({"error": "Image too large. Max 8 MB."}, status=400)

        product_name_hint = request.POST.get("product_name_hint", "").strip()
        try:
            model_index = int(request.POST.get("model_index", 0))
            batch_size = int(request.POST.get("batch_size", 1))
        except ValueError:
            model_index = 0
            batch_size = 1

        # ── Call Claude Vision ────────────────────────────────────────────────
        image_bytes = image_file.read()

        try:
            seo = generate_seo_content(
                image_file=io.BytesIO(image_bytes),
                product_name=product_name_hint or "this product",
                image_media_type=media_type,
                model_index=model_index,
                batch_size=batch_size,
            )
        except ClaudeServiceError as e:
            logger.error("Claude error during generate: %s", e)
            return JsonResponse({"error": str(e)}, status=502)

        # ── Cache image + SEO for the submit step ─────────────────────────────
        token = uuid.uuid4().hex
        logger.info("[Token %s] Caching image (%d bytes) with media type %s", token, len(image_bytes), media_type)
        cache_image(token, image_bytes, media_type)
        
        cache.set(
            SEO_CACHE_KEY.format(token=token),
            seo.to_dict(),
            timeout=SEO_CACHE_TTL,
        )
        
        # Immediate verification
        verification = cache.get(f"seo_img_{token}")
        if verification:
             logger.info("[Token %s] Cache verification SUCCESSFUL — image found in cache.", token)
        else:
             logger.error("[Token %s] Cache verification FAILED — image missing immediately after set!", token)

        response_data = seo.to_dict()
        response_data["token"] = token
        return JsonResponse(response_data)


class SubmitView(View):
    """
    POST /submit/

    Receives:
      token        — session token linking to cached image + SEO
      partner_name — partner's registered name on thotfy.com
      product_name — product to update
      email        — for notifications

    Queues the Celery pipeline, returns { task_id }.
    """

    def post(self, request):
        token        = request.POST.get("token", "").strip()
        partner_name = request.POST.get("partner_name", "").strip()
        product_name = request.POST.get("product_name", "").strip()
        email        = request.POST.get("email", "").strip()

        # ── Basic validation ──────────────────────────────────────────────────
        errors = {}
        if not token:
            errors["token"] = "Session expired. Please re-upload your image."
        if not partner_name:
            errors["partner_name"] = "Partner name is required."
        if not product_name:
            errors["product_name"] = "Product name is required."
        if not email or "@" not in email:
            errors["email"] = "A valid email address is required."
        if errors:
            return JsonResponse({"errors": errors}, status=400)

        # ── Verify cached image still exists ─────────────────────────────────
        seo_data = cache.get(SEO_CACHE_KEY.format(token=token))
        if not seo_data:
            return JsonResponse(
                {"errors": {"token": "Session expired — please re-upload your image."}},
                status=400,
            )

        # ── Queue pipeline — use token as the image cache key ─────────────────
        task = run_seo_pipeline.apply_async(
            args=[token, partner_name, product_name, email],
            task_id=token,   # reuse token as task id so the image is already keyed correctly
        )

        logger.info(
            "Pipeline queued — task=%s partner=%s product=%s email=%s",
            task.id, partner_name, product_name, email,
        )

        return JsonResponse({"task_id": task.id})


class StatusView(View):
    """
    GET /status/<task_id>/

    Returns current Celery task state + result fields once done.
    Polled every 3 s by the frontend.
    """

    def get(self, request, task_id: str):
        from celery.result import AsyncResult
        result = AsyncResult(task_id)

        state = result.state   # PENDING / STARTED / SUCCESS / FAILURE / RETRY

        if state == "SUCCESS":
            data = result.result or {}
            return JsonResponse({
                "status": data.get("status", "success"),
                "seo_title":        data.get("seo_title", ""),
                "meta_description": data.get("meta_description", ""),
                "admin_url":        data.get("admin_url", ""),
            })

        if state == "FAILURE":
            return JsonResponse({
                "status": "failed",
                "error": str(result.result),
            })

        # PENDING / STARTED / RETRY → still processing
        return JsonResponse({"status": "processing"})

class DownloadDocxView(View):
    """
    GET /download/<task_id>/
    Generates a Word Document (.docx) for the given task result.
    """
    def get(self, request, task_id: str):
        from celery.result import AsyncResult
        result = AsyncResult(task_id)

        # Allow downloading if the record is in cache OR in successful Celery result
        data = cache.get(f"seo_result_{task_id}")
        if not data and result.state == "SUCCESS":
            data = result.result or {}

        if not data:
            return HttpResponse("Result not found or task not completed.", status=404)

        seo_title = data.get("seo_title", "Untitled Product")
        slug = data.get("slug", "")
        meta_title = data.get("meta_title", "")
        meta_description = data.get("meta_description", "")
        seo_description = data.get("seo_description", "")
        captions = data.get("captions", [])
        keywords = data.get("keyword_analysis", "")
        competitors = data.get("competitor_analysis", "")

        # Create Word Doc
        doc = docx.Document()
        doc.add_heading(f"SEO Strategy: {seo_title}", 0)

        # ── Section 1: Metadata ─────────────────
        doc.add_heading("1. Product Metadata", level=1)
        doc.add_paragraph(f"SEO Title: {seo_title}")
        doc.add_paragraph(f"Slug: {slug}")
        doc.add_paragraph(f"Meta Title: {meta_title}")
        doc.add_paragraph(f"Meta Description: {meta_description}")

        # ── Section 2: Description ──────────────
        doc.add_heading("2. Product Description", level=1)
        doc.add_paragraph(seo_description)

        # ── Section 3: Social Captions ──────────
        if captions:
            doc.add_heading("3. Social Media Captions (5 Post Ideas)", level=1)
            for i, cap in enumerate(captions, 1):
                doc.add_paragraph(f"Post {i}: {cap}")

        # ── Section 4: Deep Analysis ─────────────
        doc.add_heading("4. Keyword Analysis", level=1)
        doc.add_paragraph(keywords)

        doc.add_heading("5. Competitor Analysis & Positioning", level=1)
        doc.add_paragraph(competitors)

        # Output to stream
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)

        response = HttpResponse(
            f.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="SEO_{task_id[:8]}.docx"'
        return response


@method_decorator(login_required, name='dispatch')
class HubOverviewView(View):
    """
    GET /hub/

    Render the SEO Operations Hub overview page that consolidates platform modules,
    strategy workflows, implementation phases, architecture, and success metrics.
    """

    def get(self, request):
        overview = get_hub_overview()
        return render(request, "seo_tool/hub_overview.html", {"hub": overview})


@method_decorator(login_required, name='dispatch')
class HubOverviewApiView(View):
    """
    GET /api/hub/overview/

    Return machine-readable SEO Operations Hub overview data for frontend/API clients.
    """

    def get(self, request):
        return JsonResponse(get_hub_overview())
